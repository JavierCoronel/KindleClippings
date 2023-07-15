from __future__ import print_function
import re
import io
import os
import argparse
import datetime
from typing import List
from fpdf import FPDF
import docx
import pandas as pd
import Levenshtein

def remove_chars(s, end_directory=""):
    """
    This is a utility function that removes special characters from the string, so that it can
    become a valid filename.
    :param s: input string
    :return: the input string, stripped of special characters
    """
    # Replace colons with a hyphen so "A: B" becomes "A - B"
    s = re.sub(" *: *", " - ", s)
    # Remove question marks or ampersands
    s = s.replace("?", "").replace("&", "and")
    # Replace ( ) with a hyphen so "this (text)" becomes "this - text"
    s = re.sub(r"\((.+?)\)", r"- \1", s)
    # Delete filename chars tht are not alphanumeric or ; , _ -
    s = re.sub(r"[^a-zA-Z\d\s\w;,_-]+", "", s)
    # Trim off anything that isn't a word at the start & end
    s = re.sub(r"^\W+|\W+$", "", s)

    max_length = 245 - len(end_directory)  # max file size limited to 255.
    s = s[:max_length]
    return s


def insert_line_break_in_pdf(pdf_file: FPDF, num_breaks: int = 1) -> FPDF:
    """
    Inserts a line break in a pdf for num_breaks times
    """
    while num_breaks != 0:
        pdf_file.multi_cell(0, 5, "", 0)
        num_breaks -= 1

    return pdf_file


def insert_bar_separator_in_pdf(pdf_file: FPDF):
    """
    Inserts a bar separator in a pdf, useful to separate highlights
    """
    pdf_file = insert_line_break_in_pdf(pdf_file)
    pdf_file.set_draw_color(191, 191, 191)
    pdf_file.line(40, pdf_file.y, 150, pdf_file.y)
    pdf_file = insert_line_break_in_pdf(pdf_file)

    return pdf_file


def extract_date_and_time(string):
    """Extracts the date and time from the string and returns a datetime object."""
    date_and_time_str = string.split("Added on ")[1]
    _, date_time_str = date_and_time_str.split(", ")
    date_time = datetime.datetime.strptime(date_time_str, "%d %B %Y %H:%M:%S")

    return date_time


def calculate_similarity(string1, string2):
    if string1 is None or string2 is None:
        return 0
    distance = Levenshtein.distance(string1, string2)
    similarity = 1 - (distance / max(len(string1), len(string2)))
    return similarity


def filter_repeated_highlights(df_highlights: pd.DataFrame)-> pd.DataFrame:
    """Takes a dataframe with highlights and removes the repetaed highlights. The most recent highlight is kept.

    Parameters
    ----------
    df_highlights : pd.DataFrame
        Dataframe with the list of highlights and the columns ["highlight", "meta", "date"]
    """
    df_highlights["next_highlight"] = df_highlights["highlight"].shift(-1)
    df_highlights["correlaton"] = df_highlights.apply(lambda x: calculate_similarity(x["highlight"], x["next_highlight"]), axis=1)
    df_highlights["time_diff"] = df_highlights["date"].shift(-1) - df_highlights["date"]
    df_highlights["correlaton_greater"] = df_highlights["correlaton"] > 0.3
    df_highlights["time_diff_greater"] = df_highlights["time_diff"] < datetime.timedelta(minutes=1)
    df_highlights["to_remove"] = df_highlights["correlaton_greater"] & df_highlights["time_diff_greater"]
    df_highlights =  df_highlights[~df_highlights["to_remove"]]
    df_highlights.reset_index(drop=True, inplace=True)
    return df_highlights[["highlight", "meta", "date"]]

def create_df_from_highlights_list(higlights_list: List)-> pd.DataFrame:
    """Takes a list of highlights and creates a dataframe with the columns ["highlight", "meta", "date"]

    Parameters
    ----------
    higlights_list : List
        List containing the text of the highlights
    """
    filtered_highlights = [string for string in higlights_list if string != "..." and string != ""]

    meta_regex_pattern = r"(Your.*\| Added on)"

    highlights_meta = []
    highlights_date = []
    highlights_text = []
    for highlight_line in filtered_highlights:
        if re.search(meta_regex_pattern, highlight_line):
            highlights_meta.append(highlight_line)
            highlight_date = extract_date_and_time(highlight_line)
            highlights_date.append(highlight_date)
        else:
            highlights_text.append(highlight_line)
    df_highlights = pd.DataFrame({'highlight': highlights_text, 'meta': highlights_meta, 'date': highlights_date})
    df_highlights = filter_repeated_highlights(df_highlights)

    return df_highlights

def add_page_number(pdf_file: FPDF, page_number:int)->FPDF:
    """Adds a page number to a new page in a FPDF

    Parameters
    ----------
    pdf_file : FPDF
        File containing the pages of the pdf
    page_number : int
        Number of page to add
    """
    pdf_file.set_y(20)
    pdf_file.set_font("lisboa", "", 11)
    pdf_file.set_text_color(77, 77, 77)
    pdf_file.multi_cell(0, 10, f" - Page {page_number} - ", align="C")  # Add the page number to the new page
    return pdf_file


def prepare_pdf_document(highlights: str, include_clip_meta=False, title: str = "Your Notes And Highlights") -> FPDF:
    """
    Will create pdf document from the notes

    :param highlights:
    :return: FPDF
    """
    df_highlights = create_df_from_highlights_list(highlights)

    pdf_file = FPDF()
    pdf_file.add_page()
    pdf_file.add_font("lisboa", "", "media/Lisboa.ttf", uni=True)
    pdf_file.set_font("lisboa", "", 22)
    pdf_file.set_margins(25, 40, 25)
    pdf_file = insert_line_break_in_pdf(pdf_file, 3)
    pdf_file.multi_cell(0, 8, title, align="C")
    pdf_file = insert_line_break_in_pdf(pdf_file, 2)

    page_number = 1
    for _, highlight_content in df_highlights.iterrows():
        preivous_height = pdf_file.y
        # create muti-cell pdf object and add text to it
        pdf_file.set_font("lisboa", "", 15)
        pdf_file.set_text_color(0, 0, 0)
        pdf_file.multi_cell(0, 7, highlight_content["highlight"], 0)

        pdf_file.set_font("lisboa", "", 11)
        pdf_file.set_text_color(77, 77, 77)
        pdf_file.multi_cell(0, 5, highlight_content["meta"], 0)
        pdf_file = insert_bar_separator_in_pdf(pdf_file)

        new_height = pdf_file.y
        if new_height < preivous_height:  # Check if the current content exceeds the page height
            page_number += 1
            pdf_file = add_page_number(pdf_file, page_number)
            pdf_file.set_y(new_height)

    return pdf_file


def convert_to_format(path, file_name, format, include_clip_meta=False):
    """
    Will get text file and will convert to specified output

    :param path:
    :param file_name:
    :param format:
    :return: name of the file created
    """
    output_file_name = file_name[0:-4] + "." + format
    with open(path + file_name, "r+", encoding="utf8") as txt_file:

        paragraph = txt_file.read().split("\n")
        if format == "pdf":
            pdf_file = prepare_pdf_document(paragraph, include_clip_meta, file_name[:-4])
            pdf_file.output(path + output_file_name)

        elif format == "docx":
            docx_file = docx.Document()
            docx_file.add_heading(file_name[0:-4], 0)

            for para in paragraph:
                # add a paragraph and store the object in a variable
                docx_file.add_paragraph(para)
            docx_file.save(path + output_file_name)

    return output_file_name


def create_file_by_type(end_directory, format, include_clip_meta=False):
    """
    Will iterate over all text files and will convert and create file with specified format
    Currently Only pdf and docx are supported

    :param end_directory:
    :param format:
    :return: list of output filenames
    """
    output_files = []

    # get files in and directory
    files = [f for f in os.listdir(end_directory) if os.path.isfile(end_directory + f)]

    for file in files:
        if file[-3:] == "txt":
            output_files.append(convert_to_format(end_directory, file, format, include_clip_meta))

    return output_files


def parse_clippings(source_file, end_directory, encoding="utf-8", format="txt", include_clip_meta=False):
    """
    Each clipping always consists of 5 lines:
    - title line
    - clipping info/metadata
    - a blank line
    - clipping text
    - a divider made up of equals signs
    Thus we can parse the clippings, and organise them by book.

    :param end_directory: the output directory where all of organised highlights will go
    :type end_directory: str
    :return: organises kindle highlights by book .
    """

    # Check that the source file (on the kindle) exists
    if not os.path.isfile(source_file):
        raise IOError("ERROR: cannot find " + source_file)

    # Create the output directory if it doesn't exist
    if not os.path.exists(end_directory):
        os.makedirs(end_directory)

    # This will keep track of the titles that we have already processed
    output_files = set()
    title = ""

    # Open clippings textfile and read data in lines
    with io.open(source_file, "r", encoding=encoding, errors="ignore") as f:
        # Individual highlights within clippings are separated by ==========
        for highlight in f.read().split("=========="):
            # For each highlight, we split it into the lines
            lines = highlight.split("\n")[1:]
            # Don't try to write if we have no body
            if len(lines) < 3 or lines[3] == "":
                continue
            # Set title and trim the hex character
            title = lines[0]
            if title[0] == "\ufeff":
                title = title[1:]

            # Remove characters and create path
            outfile_name = remove_chars(title, end_directory) + ".txt"
            path = end_directory + "/" + outfile_name

            # If we haven't seen title yet, set mode to write. Else, set to append.
            if outfile_name not in (list(output_files) + os.listdir(end_directory)):
                mode = "w"
                output_files.add(outfile_name)
                current_text = ""
            else:
                # If the title exists, read it as text so that we won't append duplicates
                mode = "a"
                with io.open(path, "r", encoding=encoding, errors="ignore") as textfile:
                    current_text = textfile.read()

            clipping_text = lines[3]
            clip_meta = lines[1]

            with io.open(path, mode, encoding=encoding, errors="ignore") as outfile:
                # Write out the the clippings text if it's not already there
                if clipping_text not in current_text:
                    outfile.write(clipping_text + "\n")
                    if include_clip_meta:
                        outfile.write(clip_meta + "\n")
                    outfile.write("\n...\n\n")

    # create additional file based on format
    if format in ["pdf", "docx"]:
        formatted_out_files = create_file_by_type(end_directory, format, include_clip_meta)
        output_files.update(formatted_out_files)
    else:
        print("Invalid format mentioned. Only txt file will be created")
        args.format = "txt"

    print("\nExported titles:\n")
    for i in output_files:
        print(i)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Extract kindle clippings into a folder with nice text files")
    parser.add_argument("-source", type=str, default="/Volumes/Kindle")
    parser.add_argument("-destination", type=str, default="./")
    parser.add_argument("-encoding", type=str, default="utf8")
    parser.add_argument("-format", type=str, default="txt")
    parser.add_argument("-include_clip_meta", type=bool, default=False)
    args = parser.parse_args()

    if args.source[-4:] == ".txt":
        source_file = args.source
    elif args.source[-1] == "/":
        source_file = args.source + "/My Clippings.txt"
    else:
        source_file = args.source + "/My Clippings.txt"

    if args.destination[-1] == "/":
        destination = args.destination + "KindleClippings/"
    else:
        destination = args.destination + "/KindleClippings/"

    parse_clippings(source_file, destination, args.encoding, args.format, args.include_clip_meta)

