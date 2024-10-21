from pathlib import Path
import logging
import docx
from .pdf_helpers import prepare_pdf_document

logging.basicConfig(format='%(asctime)s %(name)s %(message)s', level=logging.DEBUG)

logger = logging.getLogger(__name__)

def convert_to_format(file_path: str, output_format: str) -> str:
    """
    Converts a text file into the specified format (pdf or docx).

    Parameters
    ----------
    file_path : str
        Path of the text file to convert.
    output_format : str
        Output format, either 'pdf' or 'docx'.
    
    Returns
    -------
    str
        Path of the created file.
    """
    formatted_file_path = f"{Path(file_path).parent}/{Path(file_path).stem}.{output_format}"
    file_name = Path(file_path).stem

    with open(file_path, "r+", encoding="utf8") as txt_file:
        paragraph = txt_file.read().split("\n")

        if output_format == "pdf":
            pdf_file = prepare_pdf_document(paragraph, file_name)
            pdf_file.output(formatted_file_path)

        elif output_format == "docx":
            docx_file = docx.Document()
            docx_file.add_heading(file_name, 0)

            for para in paragraph:
                docx_file.add_paragraph(para)
            docx_file.save(formatted_file_path)

    return formatted_file_path


def convert_files_to_formatted_highlights(files: list[str], output_format: str) -> list[str]:
    """
    Iterates over list of text filesand converts them to the specified format.

    Parameters
    ----------
    files : list[str]
        Path list of text files to convert.
    output_format : str
        Output format, either 'pdf' or 'docx'.
    
    Returns
    -------
    List[str]
        List of paths to created files.
    """
    output_files = []
    logger.info(f'Converting highlights to {output_format} for {len(files)} books')
    for file in files:
        logger.info(f'Converting book {Path(file).name}')
        converted_file = convert_to_format(file, output_format)
        output_files.append(converted_file)

    return output_files
